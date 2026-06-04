"""
Generación de informe Word (.docx) para Llamadas SE.
Incluye espacio reservado para pegar captura del mapa web (mejor calidad que PDF).
"""
from __future__ import annotations

import io
from typing import Any

from app.blueprints.analisis_llamadas_se.pdf_informe import (
    DEFAULT_COLUMNS,
    INFORME_COLUMN_DEFS,
    _kpis,
    _now_local_str,
)
from app.models.analisis_llamadas_se import LlamadaSE

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
except Exception:
    Document = None


def _set_cell_shading(cell, fill_hex: str):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _set_row_height(row, height_cm: float):
    tr = row._tr
    tr_pr = tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(int(height_cm * 567)))  # twips (~1/567 cm)
    tr_height.set(qn("w:hRule"), "atLeast")
    tr_pr.append(tr_height)


def build_informe_docx(
    rows: list[LlamadaSE],
    *,
    filters_lines: list[str],
    columns: list[str],
    include_cover: bool = True,
    include_map: bool = True,
    include_table: bool = True,
    unidad_nombre: str = "",
    usuario_nombre: str = "",
) -> bytes:
    if Document is None:
        raise RuntimeError("Falta dependencia python-docx en el servidor.")

    valid_cols = [c for c in columns if c in INFORME_COLUMN_DEFS]
    if include_table and not valid_cols:
        valid_cols = DEFAULT_COLUMNS.copy()

    doc = Document()
    stats = _kpis(rows)
    now_txt = _now_local_str()

    if include_cover:
        title = doc.add_heading("Informe - Llamadas SE", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for line in (
            f"Generado: {now_txt}",
            f"Unidad: {unidad_nombre or '—'}",
            f"Usuario: {usuario_nombre or '—'}",
        ):
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        doc.add_paragraph()
        if filters_lines:
            for fl in filters_lines:
                p = doc.add_paragraph(fl)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p = doc.add_paragraph("Sin filtros adicionales (dataset completo de la unidad)")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        kpi = doc.add_table(rows=2, cols=3)
        kpi.style = "Table Grid"
        headers = ["Total llamadas", "Alerta venta", "Alerta consumo"]
        values = [str(stats["total"]), str(stats["venta"]), str(stats["consumo"])]
        for i, h in enumerate(headers):
            cell = kpi.rows[0].cells[i]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            _set_cell_shading(cell, "198754")
            kpi.rows[1].cells[i].text = values[i]
            kpi.rows[1].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

    if include_map:
        h = doc.add_heading("MAPA", level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info = doc.add_paragraph(
            f"Puntos en mapa: {stats['con_coords']} "
            f"(llamadas sin coordenadas: {stats['total'] - stats['con_coords']})"
        )
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER

        placeholder = doc.add_table(rows=1, cols=1)
        placeholder.style = "Table Grid"
        cell = placeholder.rows[0].cells[0]
        _set_row_height(placeholder.rows[0], 11.0)
        _set_cell_shading(cell, "EEF2F7")
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run("\nPEGAR AQUÍ LA CAPTURA DEL MAPA\n")
        r1.bold = True
        r1.font.size = Pt(14)
        r1.font.color.rgb = RGBColor(73, 80, 87)

        steps = [
            "1. En SIOC, abrí Mapa con los mismos filtros de este informe.",
            "2. Capturá la pantalla del mapa (Win + Shift + S, o Impr Pant).",
            "3. Hacé clic en este recuadro y pegá la imagen (Ctrl + V).",
        ]
        for step in steps:
            ps = cell.add_paragraph(step)
            ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in ps.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(108, 117, 125)

        leg = doc.add_paragraph(
            "Leyenda: rojo = venta, verde = consumo, amarillo = sospecha, azul = otras alertas."
        )
        leg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in leg.runs:
            run.font.size = Pt(9)
            run.font.italic = True

        doc.add_page_break()

    if include_table and valid_cols:
        ht = doc.add_heading(f"Detalle de llamadas ({stats['total']} registros)", level=1)
        ht.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tbl = doc.add_table(rows=1, cols=len(valid_cols))
        tbl.style = "Table Grid"
        getters = [INFORME_COLUMN_DEFS[c][1] for c in valid_cols]
        for i, c in enumerate(valid_cols):
            cell = tbl.rows[0].cells[i]
            cell.text = INFORME_COLUMN_DEFS[c][0]
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(8)
            _set_cell_shading(cell, "0D6EFD")
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

        for row in rows:
            cells = tbl.add_row().cells
            for i, fn in enumerate(getters):
                txt = fn(row) or "—"
                cells[i].text = str(txt)[:500]
                for run in cells[i].paragraphs[0].runs:
                    run.font.size = Pt(7)

    if not include_cover and not include_map and not include_table:
        doc.add_paragraph("Informe vacío: no se seleccionó ninguna sección.")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
